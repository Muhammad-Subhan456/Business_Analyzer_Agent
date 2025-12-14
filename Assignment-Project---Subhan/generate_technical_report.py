"""
Generate Technical Report as Word Document
Creates a comprehensive 10-15 page technical report with all required sections.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    new_run.append(rPr)
    new_run.text = text
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink


def create_technical_report():
    """Create comprehensive technical report."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title Page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('AI Business Analyst Agent')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(20, 184, 166)
    
    doc.add_paragraph()
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run('Technical Report')
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    from datetime import datetime
    date_para.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    
    doc.add_page_break()
    
    # Table of Contents
    toc_heading = doc.add_heading('Table of Contents', 1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    toc_items = [
        ('1. Introduction & Problem Statement', 1),
        ('2. System Architecture', 2),
        ('3. Agent Design & Reasoning Logic', 3),
        ('4. Dataset Description', 4),
        ('5. Algorithmic/LLM Methods Used', 5),
        ('6. Pydantic Models & Validation Strategy', 6),
        ('7. Database Schema & Logging Approach', 7),
        ('8. UI Design', 8),
        ('9. Testing & Evaluation', 9),
        ('10. Challenges & Limitations', 10),
        ('11. Conclusion & Future Enhancements', 11),
        ('12. Technologies Used', 12),
    ]
    
    for item, page_num in toc_items:
        para = doc.add_paragraph()
        para.add_run(item).font.size = Pt(11)
        para.add_run(f' ................. {page_num}').font.size = Pt(11)
        para.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_page_break()
    
    # Section 1: Introduction & Problem Statement
    doc.add_heading('1. Introduction & Problem Statement', 1)
    
    doc.add_paragraph(
        'Business analysis is a critical process for investors, financial analysts, and decision-makers who need comprehensive insights into companies and their market positions. Traditional business analysis requires significant time investment, manual data collection from multiple sources, and expertise in financial analysis, competitive intelligence, and market research.'
    )
    
    doc.add_paragraph(
        'The problem addressed by this project is the lack of an automated, intelligent system that can perform comprehensive business analysis efficiently. Manual analysis processes are time-consuming, prone to human error, and require domain expertise across multiple areas including finance, market research, and competitive analysis.'
    )
    
    doc.add_paragraph(
        'This project presents an AI-powered Business Analyst Agent system that automates the entire business analysis workflow. The system leverages multi-agent artificial intelligence to fetch real-time financial data, research competitors, gather market news, perform financial analysis, and generate comprehensive business reports automatically.'
    )
    
    doc.add_paragraph(
        'The solution addresses key challenges:'
    )
    
    bullet_points = [
        'Automated data collection from multiple sources including stock markets, web searches, and financial databases',
        'Intelligent analysis using specialized AI agents for different aspects of business analysis',
        'Comprehensive report generation that synthesizes financial data, competitive landscape, and market insights',
        'Real-time data processing with up-to-date information',
        'User-friendly interface that makes complex analysis accessible to non-experts',
        'Structured data storage and validation ensuring quality and traceability'
    ]
    
    for point in bullet_points:
        para = doc.add_paragraph(point, style='List Bullet')
    
    doc.add_paragraph(
        'The system is designed to reduce analysis time from hours or days to minutes while maintaining high quality and comprehensive coverage of all relevant business aspects.'
    )
    
    doc.add_page_break()
    
    # Section 2: System Architecture
    doc.add_heading('2. System Architecture', 1)
    
    doc.add_heading('2.1 Overall Architecture', 2)
    
    doc.add_paragraph(
        'The AI Business Analyst Agent follows a multi-layered architecture with clear separation of concerns. The system is built using a modular design pattern that enables scalability, maintainability, and extensibility.'
    )
    
    doc.add_paragraph(
        'The architecture consists of four main layers:'
    )
    
    layers = [
        ('Presentation Layer', 'Streamlit-based web interface that provides user interaction, input collection, and report display. This layer handles all user-facing operations and provides real-time feedback during analysis.'),
        ('Orchestration Layer', 'CrewAI framework that coordinates multiple specialized agents. This layer manages task sequencing, agent communication, and workflow execution. The BusinessAnalystCrew class serves as the central orchestrator.'),
        ('Agent Layer', 'Specialized AI agents divided into two categories: Tool Agents for data retrieval and Reasoning Agents for analysis. Each agent has a specific role and set of capabilities.'),
        ('Data Layer', 'SQLite database for persistent storage, external APIs for data fetching (yfinance, Serper), and validation models for data quality assurance.')
    ]
    
    for layer_name, description in layers:
        para = doc.add_paragraph()
        para.add_run(f'{layer_name}: ').bold = True
        para.add_run(description)
    
    doc.add_heading('2.2 Data Flow', 2)
    
    doc.add_paragraph(
        'The system follows a sequential data flow pattern where each stage depends on the completion of previous stages:'
    )
    
    flow_steps = [
        'User submits a stock ticker symbol through the Streamlit interface',
        'BusinessAnalystCrew creates a query record in the database and initializes the workflow',
        'Tool Agents execute data gathering tasks: Stock Data Agent fetches financial data, Web Search Agent finds competitors and news',
        'Reasoning Agents process the gathered data: Financial Analyst analyzes financial metrics, Competitor Analyst evaluates competitive landscape',
        'Report Writer Agent synthesizes all analysis into a comprehensive business report',
        'Report is validated using Pydantic models to ensure quality and completeness',
        'Validated report and metadata are stored in the database',
        'Report is displayed to the user and made available for PDF download'
    ]
    
    for i, step in enumerate(flow_steps, 1):
        para = doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('2.3 Component Interaction', 2)
    
    doc.add_paragraph(
        'Components interact through well-defined interfaces. The CrewAI framework manages agent communication through shared context tasks. Agents receive input from previous tasks and produce output that becomes input for subsequent tasks. The database layer provides persistence and logging capabilities throughout the process.'
    )
    
    doc.add_page_break()
    
    # Section 3: Agent Design & Reasoning Logic
    doc.add_heading('3. Agent Design & Reasoning Logic', 1)
    
    doc.add_heading('3.1 Agent Classification', 2)
    
    doc.add_paragraph(
        'The system employs a two-tier agent architecture that separates data retrieval from analysis and reasoning:'
    )
    
    doc.add_heading('3.1.1 Tool Agents', 3)
    
    doc.add_paragraph(
        'Tool Agents are specialized for data retrieval and execution of specific actions. They do not perform reasoning or analysis but focus on accurate data collection.'
    )
    
    tool_agents = [
        ('Stock Data Agent', 'Retrieves financial data using yfinance library. Fetches historical stock prices, company information, financial metrics, and market data. Uses YFinanceStockTool and YFinanceCompanyInfoTool to gather comprehensive financial information.'),
        ('Web Search Agent', 'Performs web searches using Serper API to find competitors, market news, and relevant business information. Uses SerperDevTool to execute search queries and return results with URLs and snippets.'),
        ('Web Scraper Agent', 'Extracts content from specific URLs when detailed information is needed. Uses ScrapeWebsiteTool to retrieve web page content and TextCleanerTool to clean and normalize the extracted text.')
    ]
    
    for agent_name, description in tool_agents:
        para = doc.add_paragraph()
        para.add_run(f'{agent_name}: ').bold = True
        para.add_run(description)
    
    doc.add_heading('3.1.2 Reasoning Agents', 3)
    
    doc.add_paragraph(
        'Reasoning Agents use Large Language Models to analyze data, make decisions, and generate insights. They process information from Tool Agents and produce analytical outputs.'
    )
    
    reasoning_agents = [
        ('Financial Analyst Agent', 'Analyzes financial data to provide insights on company valuation, financial health, growth trajectory, and investment potential. Processes stock price trends, financial ratios, profitability metrics, and growth rates. Produces comprehensive financial analysis with specific metrics and recommendations.'),
        ('Competitor Analyst Agent', 'Analyzes competitive landscape by identifying key competitors, comparing market positions, evaluating competitive advantages, and identifying market threats. Creates comparison tables and provides strategic insights on market dynamics.'),
        ('Report Writer Agent', 'Synthesizes all analysis components into a comprehensive, well-structured business report. Creates executive summaries, structures content into clear sections, formats for readability, and highlights key takeaways. Ensures professional presentation suitable for business decision-makers.')
    ]
    
    for agent_name, description in reasoning_agents:
        para = doc.add_paragraph()
        para.add_run(f'{agent_name}: ').bold = True
        para.add_run(description)
    
    doc.add_heading('3.2 Reasoning Logic', 2)
    
    doc.add_paragraph(
        'Each Reasoning Agent follows a structured reasoning process:'
    )
    
    reasoning_steps = [
        'Context Understanding: Agent receives context from previous tasks, including raw data and preliminary analysis',
        'Data Processing: Agent processes the input data, identifying key patterns, metrics, and relationships',
        'Analysis Execution: Agent applies domain knowledge and analytical frameworks to evaluate the data',
        'Insight Generation: Agent generates insights, conclusions, and recommendations based on the analysis',
        'Output Formatting: Agent structures the output according to task requirements and expected format'
    ]
    
    for i, step in enumerate(reasoning_steps, 1):
        para = doc.add_paragraph(f'{i}. {step}')
    
    doc.add_paragraph(
        'The reasoning process is guided by agent roles, goals, and backstories defined in the CrewAI framework. Each agent has a specific expertise area and applies relevant analytical frameworks to produce high-quality outputs.'
    )
    
    doc.add_heading('3.3 Agent Communication', 2)
    
    doc.add_paragraph(
        'Agents communicate through the CrewAI context mechanism. Tasks define dependencies where later tasks receive output from earlier tasks as context. This enables information flow from data gathering through analysis to final report generation.'
    )
    
    doc.add_page_break()
    
    # Section 4: Dataset Description
    doc.add_heading('4. Dataset Description', 1)
    
    doc.add_heading('4.1 Data Sources', 2)
    
    doc.add_paragraph(
        'The system integrates data from multiple sources to provide comprehensive business analysis:'
    )
    
    data_sources = [
        ('Yahoo Finance (yfinance)', 'Primary source for financial data. Provides real-time and historical stock prices, company fundamentals, financial ratios, market metrics, and company information. Data includes price history, volume, market cap, P/E ratios, revenue, earnings, and other key financial indicators. The library provides free access to comprehensive financial data without requiring API keys.'),
        ('Serper API', 'Provides web search capabilities for finding competitors, market news, and business information. Returns search results with titles, snippets, and URLs. Used for gathering qualitative information about companies, their competitors, and market trends. The free tier provides 2500 searches.'),
        ('Web Scraping', 'Extracts detailed content from specific web pages when needed. Used to gather comprehensive information from company websites, news articles, and other relevant sources. Content is cleaned and normalized before use.')
    ]
    
    for source_name, description in data_sources:
        para = doc.add_paragraph()
        para.add_run(f'{source_name}: ').bold = True
        para.add_run(description)
    
    doc.add_heading('4.2 Data Types', 2)
    
    doc.add_paragraph(
        'The system processes various types of data:'
    )
    
    data_types = [
        ('Quantitative Financial Data', 'Numerical data including stock prices, financial ratios, market metrics, revenue, earnings, and growth rates. Stored as structured JSON from yfinance API.'),
        ('Qualitative Business Information', 'Textual data including company descriptions, business models, product information, and market positioning. Retrieved from web searches and company information APIs.'),
        ('Competitive Intelligence', 'Information about competitors including company names, market positions, competitive advantages, and market share data. Gathered through web searches and analysis.'),
        ('Market News and Events', 'Recent news articles, earnings reports, product launches, and market-moving events. Collected through web searches and news aggregation.')
    ]
    
    for data_type, description in data_types:
        para = doc.add_paragraph()
        para.add_run(f'{data_type}: ').bold = True
        para.add_run(description)
    
    doc.add_heading('4.3 Data Processing', 2)
    
    doc.add_paragraph(
        'Data undergoes several processing steps:'
    )
    
    processing_steps = [
        'Data Retrieval: Tool Agents fetch raw data from external sources',
        'Data Cleaning: Text data is cleaned to remove noise, boilerplate, and irrelevant content',
        'Data Validation: Pydantic models validate data structure and completeness',
        'Data Analysis: Reasoning Agents process data to extract insights',
        'Data Storage: Validated data and analysis results are stored in SQLite database',
        'Data Presentation: Final reports are formatted for user consumption'
    ]
    
    for step in processing_steps:
        para = doc.add_paragraph(step, style='List Bullet')
    
    doc.add_page_break()
    
    # Section 5: Algorithmic/LLM Methods Used
    doc.add_heading('5. Algorithmic/LLM Methods Used', 1)
    
    doc.add_heading('5.1 Large Language Model', 2)
    
    doc.add_paragraph(
        'The system uses Ollama with the Llama 3.2 model as the primary LLM for all reasoning agents. Ollama provides local LLM inference, eliminating the need for external API keys and ensuring data privacy.'
    )
    
    doc.add_paragraph(
        'LLM Configuration:'
    )
    
    llm_config = [
        'Model: ollama/llama3.2',
        'Base URL: http://127.0.0.1:11434 (local Ollama server)',
        'Temperature Settings: Tool Agents use 0.1 (low for consistency), Reasoning Agents use 0.5-0.7 (medium for creativity)',
        'Max Iterations: Tool Agents (3-5), Reasoning Agents (3-5), Report Writer (3)'
    ]
    
    for config in llm_config:
        para = doc.add_paragraph(config, style='List Bullet')
    
    doc.add_heading('5.2 CrewAI Framework', 2)
    
    doc.add_paragraph(
        'CrewAI provides the multi-agent orchestration framework. Key features used:'
    )
    
    crewai_features = [
        'Agent Definition: Agents are defined with roles, goals, backstories, and tools',
        'Task Management: Tasks define what agents should do, expected outputs, and dependencies',
        'Crew Orchestration: Crew coordinates multiple agents and manages task execution',
        'Context Passing: Tasks can receive context from previous tasks, enabling information flow',
        'Sequential Processing: Tasks execute in order based on dependencies'
    ]
    
    for feature in crewai_features:
        para = doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('5.3 Analysis Algorithms', 2)
    
    doc.add_paragraph(
        'The system employs various analytical methods:'
    )
    
    algorithms = [
        ('Financial Ratio Analysis', 'Calculation and interpretation of financial ratios including P/E, P/B, PEG, profit margins, ROE, and ROA. Agents compare ratios to industry averages and historical trends.'),
        ('Trend Analysis', 'Analysis of stock price trends, revenue growth, earnings growth, and other time-series data. Identifies patterns, support/resistance levels, and momentum indicators.'),
        ('Competitive Analysis', 'Comparative analysis of companies within the same industry. Evaluates market positions, competitive advantages, and relative performance.'),
        ('SWOT Analysis', 'Structured analysis of Strengths, Weaknesses, Opportunities, and Threats. Synthesizes financial data, competitive intelligence, and market information.'),
        ('Valuation Methods', 'Assessment of company valuation using multiple metrics including P/E ratios, market cap, and growth rates. Provides fair value estimates and investment recommendations.')
    ]
    
    for algo_name, description in algorithms:
        para = doc.add_paragraph()
        para.add_run(f'{algo_name}: ').bold = True
        para.add_run(description)
    
    doc.add_heading('5.4 Prompt Engineering', 2)
    
    doc.add_paragraph(
        'Task descriptions serve as prompts that guide agent behavior. Prompts are carefully crafted to:'
    )
    
    prompt_aspects = [
        'Define clear objectives and expected outputs',
        'Specify required sections and content structure',
        'Provide context about the analysis type and scope',
        'Include formatting requirements and quality standards',
        'Guide agents to use specific tools and data sources'
    ]
    
    for aspect in prompt_aspects:
        para = doc.add_paragraph(aspect, style='List Bullet')
    
    doc.add_page_break()
    
    # Section 6: Pydantic Models & Validation Strategy
    doc.add_heading('6. Pydantic Models & Validation Strategy', 1)
    
    doc.add_heading('6.1 Validation Philosophy', 2)
    
    doc.add_paragraph(
        'The system implements comprehensive validation using Pydantic models to ensure data quality, type safety, and completeness. All major outputs pass through validation before storage or presentation.'
    )
    
    doc.add_heading('6.2 ReportValidationModel', 2)
    
    doc.add_paragraph(
        'The ReportValidationModel validates final business analysis reports:'
    )
    
    report_validation = [
        ('Required Fields', 'ticker (validated format), report_content (minimum 100 characters), report_type (Full Analysis or Quick Analysis), generated_at (timestamp)'),
        ('Type Enforcement', 'Automatic type checking and conversion. String fields validated for length and format. Timestamps validated as datetime objects.'),
        ('Completeness Scoring', 'Calculates completeness score (0.0-1.0) based on presence of required sections: Executive Summary, Company Overview, Financial Analysis, Key Takeaways. Score includes base value (0.5) plus points for each section found (0.125 each) and bonus for appropriate word count (0.1).'),
        ('Structure Scoring', 'Evaluates report structure quality (0.0-1.0) based on: number of markdown headers (0.4 for 3+ headers), bullet points/lists (0.3 for 5+ items), and data points/numbers (0.3 for 10+ numbers).'),
        ('Section Extraction', 'Automatically extracts section names from markdown headers for tracking and validation.')
    ]
    
    for aspect, description in report_validation:
        para = doc.add_paragraph()
        para.add_run(f'{aspect}: ').bold = True
        para.add_run(description)
    
    doc.add_heading('6.3 AnalysisMetadataModel', 2)
    
    doc.add_paragraph(
        'The AnalysisMetadataModel validates analysis metadata and summaries:'
    )
    
    metadata_validation = [
        ('Required Fields', 'ticker, summary (minimum 50 characters), key_decisions (minimum 20 characters), data_completeness (0.0-1.0), confidence_score (0.0-1.0)'),
        ('Score Validation', 'All scores validated to be within 0.0-1.0 range and rounded to 3 decimal places for consistency.'),
        ('Quality Calculation', 'Calculates overall quality score as weighted average: 60% data_completeness + 40% confidence_score.'),
        ('Content Validation', 'Summary and key_decisions validated for minimum length to ensure meaningful content.')
    ]
    
    for aspect, description in metadata_validation:
        para = doc.add_paragraph()
        para.add_run(f'{aspect}: ').bold = True
        para.add_run(description)
    
    doc.add_heading('6.4 Validation Workflow', 2)
    
    doc.add_paragraph(
        'Validation occurs at key points in the system:'
    )
    
    validation_points = [
        'Report Generation: After Report Writer Agent produces final report, ReportValidationModel validates structure, completeness, and content',
        'Metadata Creation: When storing analysis metadata, AnalysisMetadataModel validates scores and summaries',
        'Database Storage: Before saving to database, all data passes through validation to ensure quality',
        'Error Handling: Validation failures are logged but do not block execution, allowing graceful degradation'
    ]
    
    for point in validation_points:
        para = doc.add_paragraph(point, style='List Bullet')
    
    doc.add_page_break()
    
    # Section 7: Database Schema & Logging Approach
    doc.add_heading('7. Database Schema & Logging Approach', 1)
    
    doc.add_heading('7.1 Database Design', 2)
    
    doc.add_paragraph(
        'The system uses SQLite for persistent storage with a minimal but comprehensive schema designed for efficiency and traceability.'
    )
    
    doc.add_heading('7.2 Schema Structure', 2)
    
    doc.add_paragraph(
        'The database consists of four main tables:'
    )
    
    doc.add_heading('7.2.1 user_queries Table', 3)
    
    doc.add_paragraph(
        'Stores core query information for each analysis request. Fields include: id (primary key), ticker, company_name, analysis_type, period, status (pending/processing/completed/failed), created_at, error_message. Indexed on ticker and created_at for efficient querying.'
    )
    
    doc.add_heading('7.2.2 reports Table', 3)
    
    doc.add_paragraph(
        'Stores generated analysis reports. Fields include: id (primary key), query_id (foreign key to user_queries), ticker, report_content (full markdown report), word_count, generated_at. Indexed on query_id for fast retrieval.'
    )
    
    doc.add_heading('7.2.3 agent_logs Table', 3)
    
    doc.add_paragraph(
        'Minimal logging of agent actions. Fields include: id (primary key), query_id (foreign key), agent_name, action_summary (brief description, max 500 chars), status (success/error), timestamp. Stores summaries rather than full tool outputs to minimize storage. Indexed on query_id for efficient retrieval.'
    )
    
    doc.add_heading('7.2.4 analysis_metadata Table', 3)
    
    doc.add_paragraph(
        'Stores analysis summaries and key decisions. Fields include: id (primary key), query_id (foreign key), key_decisions (summarized agent decisions, max 1000 chars), data_completeness (0.0-1.0), confidence_score (0.0-1.0), summary (brief analysis summary, max 500 chars), created_at. Indexed on query_id.'
    )
    
    doc.add_heading('7.3 Logging Strategy', 2)
    
    doc.add_paragraph(
        'The system implements minimal but effective logging:'
    )
    
    logging_strategy = [
        'Query Tracking: Every analysis request creates a query record with status tracking',
        'Agent Actions: Key agent actions are logged with brief summaries (not full tool outputs)',
        'Error Logging: Failed operations are logged with error messages for debugging',
        'Metadata Storage: Analysis summaries and key decisions are stored for quick reference',
        'Performance: Minimal logging reduces storage overhead while maintaining traceability'
    ]
    
    for strategy in logging_strategy:
        para = doc.add_paragraph(strategy, style='List Bullet')
    
    doc.add_heading('7.4 Data Retention', 2)
    
    doc.add_paragraph(
        'The database includes cleanup functionality to remove old data. By default, data older than 90 days can be automatically removed to manage storage. The cleanup function respects foreign key relationships and deletes data in the correct order.'
    )
    
    doc.add_page_break()
    
    # Section 8: UI Design
    doc.add_heading('8. UI Design', 1)
    
    doc.add_heading('8.1 Framework and Technology', 2)
    
    doc.add_paragraph(
        'The user interface is built using Streamlit, a Python framework for creating web applications. Streamlit provides a simple, declarative API for building interactive UIs without requiring frontend development expertise.'
    )
    
    doc.add_heading('8.2 Design Philosophy', 2)
    
    doc.add_paragraph(
        'The UI follows a modern, professional design philosophy:'
    )
    
    design_principles = [
        'Dark Theme: Modern dark interface with teal accents for a professional appearance',
        'User-Friendly: Simple, intuitive interface that requires no training',
        'Real-Time Feedback: Progress indicators and status updates during analysis',
        'Responsive Layout: Adapts to different screen sizes and resolutions',
        'Accessibility: Clear typography, sufficient contrast, and logical information hierarchy'
    ]
    
    for principle in design_principles:
        para = doc.add_paragraph(principle, style='List Bullet')
    
    doc.add_heading('8.3 Interface Components', 2)
    
    doc.add_paragraph(
        'Key interface components:'
    )
    
    components = [
        ('Sidebar', 'Contains input fields for ticker symbol, company name, analysis type selection, and period selection. Also includes API key status, feature list, and database viewer for stored reports.'),
        ('Main Content Area', 'Displays analysis instructions, analysis button, progress indicators, and generated reports. Provides download functionality for PDF reports.'),
        ('Report Display', 'Shows complete reports with proper markdown formatting. Includes report header with ticker and timestamp, download button, and full report content without truncation.'),
        ('Progress Indicators', 'Visual feedback during analysis including progress bars, status messages, and stage indicators showing current analysis phase.')
    ]
    
    for component_name, description in components:
        para = doc.add_paragraph()
        para.add_run(f'{component_name}: ').bold = True
        para.add_run(description)
    
    doc.add_heading('8.4 User Experience Flow', 2)
    
    doc.add_paragraph(
        'Typical user interaction flow:'
    )
    
    ux_flow = [
        'User opens the application in a web browser',
        'User enters stock ticker symbol in the sidebar',
        'User optionally provides company name and selects analysis type and period',
        'User clicks Start Analysis button',
        'System displays progress indicators showing analysis stages',
        'Upon completion, system displays the full report',
        'User can download the report as PDF or view stored reports in the sidebar'
    ]
    
    for i, step in enumerate(ux_flow, 1):
        para = doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('8.5 Styling and Customization', 2)
    
    doc.add_paragraph(
        'The UI uses custom CSS for styling:'
    )
    
    styling_features = [
        'Custom color scheme with dark background and teal accent colors',
        'Custom fonts (Outfit for headings, JetBrains Mono for inputs)',
        'Styled cards and containers with hover effects',
        'Gradient buttons and accent elements',
        'Responsive scrollbars and layout adjustments'
    ]
    
    for feature in styling_features:
        para = doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # Section 9: Testing & Evaluation
    doc.add_heading('9. Testing & Evaluation', 1)
    
    doc.add_heading('9.1 Testing Strategy', 2)
    
    doc.add_paragraph(
        'The system includes comprehensive testing capabilities:'
    )
    
    doc.add_heading('9.1.1 Unit Testing', 3)
    
    doc.add_paragraph(
        'Individual components are tested in isolation:'
    )
    
    unit_tests = [
        'Pydantic Models: Validation models tested with valid and invalid inputs to ensure proper validation',
        'Database Operations: Database manager tested for CRUD operations, query execution, and data integrity',
        'PDF Generation: PDF generation function tested with sample reports to verify valid PDF output',
        'Tool Functions: Individual tools tested for correct data retrieval and processing'
    ]
    
    for test in unit_tests:
        para = doc.add_paragraph(test, style='List Bullet')
    
    doc.add_heading('9.1.2 Integration Testing', 3)
    
    doc.add_paragraph(
        'System components are tested together:'
    )
    
    integration_tests = [
        'End-to-End Analysis: Complete analysis workflow tested from user input to final report generation',
        'Database Integration: Validation and storage workflow tested to ensure data flows correctly',
        'Agent Communication: Agent interactions and context passing tested for correctness',
        'UI Integration: User interface tested for proper display and interaction'
    ]
    
    for test in integration_tests:
        para = doc.add_paragraph(test, style='List Bullet')
    
    doc.add_heading('9.2 Test Scripts', 2)
    
    doc.add_paragraph(
        'The project includes several test scripts:'
    )
    
    test_scripts = [
        ('test_implementation.py', 'Comprehensive test suite for database and Pydantic models. Tests validation, database operations, and integration.'),
        ('test_pdf_quick.py', 'Quick PDF generation test that creates sample reports and verifies PDF output without running full analysis. Saves testing time from 20-25 minutes to 5-10 seconds.'),
        ('view_database.py', 'Database viewer script for inspecting stored data, queries, reports, and metadata.')
    ]
    
    for script_name, description in test_scripts:
        para = doc.add_paragraph()
        para.add_run(f'{script_name}: ').bold = True
        para.add_run(description)
    
    doc.add_heading('9.3 Evaluation Metrics', 2)
    
    doc.add_paragraph(
        'System performance is evaluated using:'
    )
    
    metrics = [
        'Report Completeness: Measured by completeness_score from ReportValidationModel (target: >0.8)',
        'Report Structure: Measured by structure_score from ReportValidationModel (target: >0.7)',
        'Data Completeness: Measured by data_completeness in metadata (target: >0.8)',
        'Confidence Score: Measured by confidence_score in metadata (target: >0.75)',
        'Analysis Success Rate: Percentage of completed analyses vs failed (target: >95%)',
        'Processing Time: Time from analysis start to report generation (typical: 15-25 minutes)'
    ]
    
    for metric in metrics:
        para = doc.add_paragraph(metric, style='List Bullet')
    
    doc.add_heading('9.4 Quality Assurance', 2)
    
    doc.add_paragraph(
        'Quality is ensured through multiple mechanisms:'
    )
    
    qa_measures = [
        'Pydantic Validation: All outputs validated before storage',
        'Error Handling: Comprehensive error handling with graceful degradation',
        'Logging: Detailed logging for debugging and monitoring',
        'User Feedback: Progress indicators and status messages keep users informed',
        'Database Integrity: Foreign key constraints ensure data consistency'
    ]
    
    for measure in qa_measures:
        para = doc.add_paragraph(measure, style='List Bullet')
    
    doc.add_page_break()
    
    # Section 10: Challenges & Limitations
    doc.add_heading('10. Challenges & Limitations', 1)
    
    doc.add_heading('10.1 Technical Challenges', 2)
    
    doc.add_paragraph(
        'Several technical challenges were encountered and addressed:'
    )
    
    challenges = [
        ('Processing Time', 'Full analysis takes 15-25 minutes due to sequential processing and local LLM inference. This is a trade-off for thoroughness and privacy. Potential solutions include parallel processing and faster LLM models.'),
        ('PDF Generation', 'Initial PDF generation produced corrupted files. Resolved by using ReportLab as primary method with proper encoding and buffer handling. Multiple fallback methods ensure reliability.'),
        ('Report Display', 'Reports were initially truncated in the UI. Fixed by using direct markdown rendering and removing HTML container limitations.'),
        ('Database Integration', 'Integrating database logging without blocking analysis required careful error handling and optional database features.'),
        ('Local LLM Setup', 'Requires users to have Ollama installed and running locally. This adds setup complexity but provides privacy benefits.')
    ]
    
    for challenge_name, description in challenges:
        para = doc.add_paragraph()
        para.add_run(f'{challenge_name}: ').bold = True
        para.add_run(description)
    
    doc.add_heading('10.2 System Limitations', 2)
    
    doc.add_paragraph(
        'Current system limitations:'
    )
    
    limitations = [
        'Sequential Processing: Tasks execute one after another, limiting speed improvements',
        'Local LLM Dependency: Requires Ollama installation and local model, which may not be available on all systems',
        'Free API Limits: Serper API free tier limited to 2500 searches total',
        'Data Source Dependency: Relies on external APIs (yfinance, Serper) which may have availability issues',
        'Report Length: Reports are limited to 800-1200 words by design, which may not cover all aspects for complex companies',
        'Single Ticker Analysis: System analyzes one company at a time, not suitable for batch processing',
        'No Historical Comparison: Does not compare current analysis with previous analyses automatically'
    ]
    
    for limitation in limitations:
        para = doc.add_paragraph(limitation, style='List Bullet')
    
    doc.add_heading('10.3 Known Issues', 2)
    
    doc.add_paragraph(
        'Known issues and workarounds:'
    )
    
    known_issues = [
        ('Windows Signal Handling', 'CrewAI uses Unix signals that do not exist on Windows. Workaround: Added dummy signal definitions for Windows compatibility.'),
        ('Ollama Connection', 'If Ollama is not running, the system will fail. Workaround: Clear error messages guide users to start Ollama.'),
        ('API Key Management', 'API keys must be set as environment variables. Workaround: Clear instructions provided in README.'),
        ('Long Analysis Times', 'Users may experience long wait times. Workaround: Progress indicators and status messages provide feedback.')
    ]
    
    for issue_name, description in known_issues:
        para = doc.add_paragraph()
        para.add_run(f'{issue_name}: ').bold = True
        para.add_run(description)
    
    doc.add_page_break()
    
    # Section 11: Conclusion & Future Enhancements
    doc.add_heading('11. Conclusion & Future Enhancements', 1)
    
    doc.add_heading('11.1 Conclusion', 2)
    
    doc.add_paragraph(
        'The AI Business Analyst Agent successfully demonstrates the application of multi-agent AI systems to automate complex business analysis workflows. The system effectively combines data retrieval, intelligent analysis, and report generation into a cohesive solution that reduces analysis time while maintaining quality.'
    )
    
    doc.add_paragraph(
        'Key achievements include:'
    )
    
    achievements = [
        'Successful implementation of multi-agent architecture using CrewAI framework',
        'Integration of multiple data sources for comprehensive analysis',
        'Development of robust validation and quality assurance mechanisms',
        'Creation of user-friendly interface that makes complex analysis accessible',
        'Implementation of persistent storage and logging for traceability',
        'Production-ready PDF generation with proper formatting'
    ]
    
    for achievement in achievements:
        para = doc.add_paragraph(achievement, style='List Bullet')
    
    doc.add_paragraph(
        'The system provides a solid foundation for automated business analysis and demonstrates the potential of AI agents in financial and business intelligence applications.'
    )
    
    doc.add_heading('11.2 Future Enhancements', 2)
    
    doc.add_paragraph(
        'Potential future improvements and enhancements:'
    )
    
    enhancements = [
        ('Parallel Processing', 'Implement parallel execution of independent tasks to reduce analysis time from 15-25 minutes to 5-10 minutes. This would require restructuring the workflow to identify parallelizable tasks.'),
        ('Batch Analysis', 'Support analyzing multiple companies simultaneously. This would enable portfolio analysis and comparative studies across multiple companies.'),
        ('Historical Comparison', 'Add functionality to compare current analysis with previous analyses, tracking changes over time and identifying trends.'),
        ('Advanced Analytics', 'Implement more sophisticated financial models including DCF valuation, technical analysis indicators, and predictive modeling.'),
        ('Real-Time Updates', 'Add capability to monitor companies and send alerts when significant changes occur in financial metrics or market conditions.'),
        ('Export Formats', 'Support additional export formats including Excel, CSV for data, and PowerPoint for presentations.'),
        ('Custom Report Templates', 'Allow users to customize report structure and sections based on their specific needs.'),
        ('API Access', 'Provide REST API for programmatic access, enabling integration with other systems and automated workflows.'),
        ('Cloud Deployment', 'Deploy as a cloud service to eliminate local setup requirements and provide scalability.'),
        ('Enhanced Visualization', 'Add charts, graphs, and visualizations to reports for better data presentation.'),
        ('Multi-Language Support', 'Support analysis and reports in multiple languages for international users.'),
        ('Integration with More Data Sources', 'Add integration with additional financial data providers, news sources, and market intelligence platforms.')
    ]
    
    for enhancement_name, description in enhancements:
        para = doc.add_paragraph()
        para.add_run(f'{enhancement_name}: ').bold = True
        para.add_run(description)
    
    doc.add_heading('11.3 Long-Term Vision', 2)
    
    doc.add_paragraph(
        'The long-term vision for the system includes becoming a comprehensive business intelligence platform that provides real-time insights, predictive analytics, and strategic recommendations. The platform would serve investors, analysts, and business decision-makers with automated, intelligent analysis capabilities that scale to handle large portfolios and provide actionable insights.'
    )
    
    doc.add_page_break()
    
    # Section 12: Technologies Used
    doc.add_heading('12. Technologies Used', 1)
    
    doc.add_heading('12.1 Core Framework', 2)
    
    doc.add_paragraph(
        'CrewAI (v0.86.0+): Multi-agent orchestration framework that enables coordination of specialized AI agents. Provides agent definition, task management, and workflow orchestration capabilities.'
    )
    
    doc.add_heading('12.2 Programming Language', 2)
    
    doc.add_paragraph(
        'Python 3.10+: Primary programming language. Chosen for its extensive libraries, AI/ML ecosystem support, and ease of development.'
    )
    
    doc.add_heading('12.3 AI/ML Technologies', 2)
    
    ai_tech = [
        ('Ollama', 'Local LLM inference server. Provides privacy-preserving AI capabilities without requiring external API keys. Used with Llama 3.2 model.'),
        ('LangChain', 'Framework for building LLM applications. Used through CrewAI integration for agent communication and tool usage.'),
        ('Google Generative AI', 'Alternative LLM provider (optional). Can be used instead of Ollama for cloud-based inference.')
    ]
    
    for tech_name, description in ai_tech:
        para = doc.add_paragraph()
        para.add_run(f'{tech_name}: ').bold = True
        para.add_run(description)
    
    doc.add_heading('12.4 Data Sources & APIs', 2)
    
    data_tech = [
        ('yfinance', 'Python library for downloading financial data from Yahoo Finance. Provides free access to stock prices, company information, and financial metrics.'),
        ('Serper API', 'Google Search API for web search capabilities. Used for finding competitors, news, and market information. Free tier provides 2500 searches.'),
        ('BeautifulSoup4', 'HTML parsing library for web scraping. Used to extract and clean content from web pages.'),
        ('Requests', 'HTTP library for making API calls and downloading web content.')
    ]
    
    for tech_name, description in data_tech:
        para = doc.add_paragraph()
        para.add_run(f'{tech_name}: ').bold = True
        para.add_run(description)
    
    doc.add_heading('12.5 Database & Storage', 2)
    
    db_tech = [
        ('SQLite', 'Lightweight, file-based database. Used for persistent storage of queries, reports, logs, and metadata. No separate server required.'),
        ('Pydantic', 'Data validation library. Used for validating reports and metadata, ensuring type safety and data quality.')
    ]
    
    for tech_name, description in db_tech:
        para = doc.add_paragraph()
        para.add_run(f'{tech_name}: ').bold = True
        para.add_run(description)
    
    doc.add_heading('12.6 Frontend & UI', 2)
    
    frontend_tech = [
        ('Streamlit', 'Python web framework for building interactive UIs. Enables rapid development of data applications without frontend expertise.'),
        ('Custom CSS', 'Styling for modern, professional appearance with dark theme and custom color scheme.')
    ]
    
    for tech_name, description in frontend_tech:
        para = doc.add_paragraph()
        para.add_run(f'{tech_name}: ').bold = True
        para.add_run(description)
    
    doc.add_heading('12.7 PDF Generation', 2)
    
    pdf_tech = [
        ('ReportLab', 'Primary PDF generation library. Pure Python library for creating PDF documents. Provides reliable, standards-compliant PDF output.'),
        ('Markdown', 'Library for converting markdown to HTML. Used as intermediate step in PDF generation process.'),
        ('WeasyPrint', 'Alternative PDF generator (optional). HTML/CSS to PDF converter.'),
        ('xhtml2pdf', 'Fallback PDF generator. HTML to PDF converter for compatibility.')
    ]
    
    for tech_name, description in pdf_tech:
        para = doc.add_paragraph()
        para.add_run(f'{tech_name}: ').bold = True
        para.add_run(description)
    
    doc.add_heading('12.8 Development Tools', 2)
    
    dev_tools = [
        ('python-dotenv', 'Environment variable management. Loads API keys and configuration from .env files.'),
        ('pandas', 'Data manipulation and analysis. Used for processing financial data.'),
        ('numpy', 'Numerical computing. Used for calculations and data processing.')
    ]
    
    for tool_name, description in dev_tools:
        para = doc.add_paragraph()
        para.add_run(f'{tool_name}: ').bold = True
        para.add_run(description)
    
    doc.add_heading('12.9 System Requirements', 2)
    
    doc.add_paragraph(
        'Minimum system requirements:'
    )
    
    requirements = [
        'Python 3.10 or higher',
        'Ollama installed and running locally (for LLM inference)',
        'Internet connection (for data fetching and web searches)',
        'Minimum 4GB RAM (8GB recommended for Ollama)',
        'Windows, macOS, or Linux operating system'
    ]
    
    for req in requirements:
        para = doc.add_paragraph(req, style='List Bullet')
    
    # Save document
    output_file = 'Technical_Report_AI_Business_Analyst_Agent.docx'
    doc.save(output_file)
    print(f'Technical report generated successfully: {output_file}')
    print(f'File location: {os.path.abspath(output_file)}')
    return output_file


if __name__ == '__main__':
    print('Generating Technical Report...')
    print('This may take a few moments...')
    create_technical_report()
    print('\nReport generation complete!')
    print('You can now open the Word document.')

